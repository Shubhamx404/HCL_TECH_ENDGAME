import streamlit as st
import fitz  # PyMuPDF
import google.generativeai as genai
from sentence_transformers import SentenceTransformer
import numpy as np
import re
import os
import datetime
from io import BytesIO
from langchain_community.document_loaders import WebBaseLoader
from typing import List, Dict, Tuple
import faiss
import pickle

# --- Configuration ---
st.set_page_config(
    page_title="JurisMind",
    page_icon="",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- Custom CSS for modern UI ---
def load_css():
    st.markdown("""
    <style>
        /* Main container */
        .main {
            background-color: #fff5f7;
        }
        /* Sidebar */
        .css-1d391kg {
            background-color: #831843 !important;
            color: white !important;
        }
        /* Header */
        .css-1v3fvcr {
            background: linear-gradient(90deg, #831843 0%, #db2777 100%) !important;
            color: white;
            padding: 1rem 2rem !important;
            border-radius: 0 0 15px 15px !important;
            margin-bottom: 2rem !important;
        }
        /* Buttons */
        .stButton>button {
            border-radius: 12px !important;
            padding: 0.5rem 1.5rem !important;
            font-weight: 600 !important;
            transition: all 0.3s ease !important;
            background: #db2777 !important;
            color: white !important;
            border: none !important;
        }
        .stButton>button:hover {
            transform: translateY(-2px) !important;
            box-shadow: 0 4px 15px rgba(219, 39, 119, 0.4) !important;
            background: #be185d !important;
        }
        /* Cards */
        .card {
            background: white;
            border-radius: 15px !important;
            padding: 1.5rem !important;
            box-shadow: 0 4px 15px rgba(0, 0, 0, 0.05) !important;
            margin-bottom: 1.5rem !important;
            border: 1px solid #fce7f3 !important;
        }
        /* Chat messages */
        .user-message {
            background-color: #fdf2f8 !important;
            border-radius: 15px 15px 0 15px !important;
            padding: 12px 16px !important;
            margin: 8px 0 !important;
            max-width: 80% !important;
            float: right !important;
            clear: both !important;
            border: 1px solid #fbcfe8 !important;
            color: #1f2937 !important;
        }
        .assistant-message {
            background-color: white !important;
            border-radius: 15px 15px 15px 0 !important;
            padding: 12px 16px !important;
            margin: 8px 0 !important;
            max-width: 80% !important;
            float: left !important;
            clear: both !important;
            border: 1px solid #f3e8ff !important;
            color: #1f2937 !important;
        }
        /* Input box */
        .stTextInput>div>div>input {
            border-radius: 12px !important;
            padding: 12px 20px !important;
            border: 1px solid #f3e8ff !important;
            box-shadow: 0 2px 8px rgba(0,0,0,0.05) !important;
        }
        /* Tabs */
        .stTabs [data-baseweb="tab-list"] {
            gap: 8px !important;
            border-bottom: 2px solid #f3e8ff !important;
            margin-bottom: 1.5rem !important;
        }
        .stTabs [data-baseweb="tab"] {
            border-radius: 12px 12px 0 0 !important;
            padding: 10px 25px !important;
            background: #fdf2f8 !important;
            margin: 0 2px !important;
            transition: all 0.3s ease !important;
            color: #831843 !important;
            font-weight: 500 !important;
        }
        .stTabs [aria-selected="true"] {
            background-color: #db2777 !important;
            color: white !important;
        }
        /* File Uploader */
        .stFileUploader>div>div {
            border: 2px dashed #f3e8ff !important;
            border-radius: 12px !important;
            padding: 2rem !important;
            background: white !important;
            transition: all 0.3s ease !important;
        }
        .stFileUploader>div>div:hover {
            border-color: #db2777 !important;
        }
    </style>
    """, unsafe_allow_html=True)

load_css()

# --- Classes & Functions ---

class FAISSVectorDB:
    """FAISS-based vector database for efficient similarity search."""

    def __init__(self, embedding_dim: int = 384):
        """Initialize FAISS index with given embedding dimension."""
        self.embedding_dim = embedding_dim
        self.index = faiss.IndexFlatL2(embedding_dim)  # L2 distance for similarity
        self.text_chunks = []
        self.metadata = []

    def add_documents(self, embeddings: np.ndarray, text_chunks: List[str], metadata: List[Dict] = None):
        """Add documents with their embeddings to the vector database."""
        if embeddings.shape[1] != self.embedding_dim:
            raise ValueError(f"Embedding dimension mismatch. Expected {self.embedding_dim}, got {embeddings.shape[1]}")

        # Add embeddings to FAISS index
        self.index.add(embeddings)

        # Store text chunks
        self.text_chunks.extend(text_chunks)

        # Store metadata if provided
        if metadata:
            self.metadata.extend(metadata)
        else:
            self.metadata.extend([{}] * len(text_chunks))

    def search(self, query_embedding: np.ndarray, top_k: int = 5) -> Tuple[List[str], List[float], List[int]]:
        """Search for similar documents using query embedding."""
        if len(self.text_chunks) == 0:
            return [], [], []

        # Adjust top_k if it exceeds available documents
        top_k = min(top_k, len(self.text_chunks))

        # Search in FAISS index
        distances, indices = self.index.search(query_embedding, top_k)

        # Retrieve corresponding text chunks
        results = []
        result_distances = []
        result_indices = []

        for i, idx in enumerate(indices[0]):
            if idx < len(self.text_chunks):
                results.append(self.text_chunks[idx])
                result_distances.append(float(distances[0][i]))
                result_indices.append(int(idx))

        return results, result_distances, result_indices

    def get_stats(self) -> Dict:
        """Get statistics about the vector database."""
        return {
            "total_documents": len(self.text_chunks),
            "embedding_dimension": self.embedding_dim,
            "index_size": self.index.ntotal
        }

    def clear(self):
        """Clear all data from the vector database."""
        self.index.reset()
        self.text_chunks = []
        self.metadata = []

def extract_text_from_file(file) -> str:
    """Extracts text from an uploaded file (PDF, DOCX, or TXT)."""
    text = f"\n--- Document: {file.name} ---\n"
    file_extension = file.name.split('.')[-1].lower()

    try:
        file.seek(0)  # Reset file pointer to the beginning

        if file_extension == 'pdf':
            with fitz.open(stream=file.read(), filetype="pdf") as doc:
                for page in doc:
                    text += page.get_text() + "\n"

        elif file_extension == 'docx':
            try:
                from docx import Document
                doc = Document(BytesIO(file.read()))
                for para in doc.paragraphs:
                    text += para.text + "\n"
            except ImportError:
                st.warning("python-docx package is required for DOCX files. Installing it now...")
                import subprocess
                import sys
                subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
                from docx import Document
                doc = Document(BytesIO(file.read()))
                for para in doc.paragraphs:
                    text += para.text + "\n"

        elif file_extension == 'txt':
            text += file.read().decode('utf-8')

        else:
            st.warning(f"Unsupported file type: {file_extension}")
            return ""

    except Exception as e:
        st.error(f"Error reading {file.name}: {e}")
        return ""

    return text + "\n"

def extract_text_from_url(url: str) -> str:
    """Extracts text from a given URL."""
    try:
        loader = WebBaseLoader(url)
        docs = loader.load()
        if not docs:
            st.warning("Could not load any content from the URL.")
            return ""
        return docs[0].page_content
    except Exception as e:
        st.error(f"Error fetching or parsing URL: {e}")
        return ""

def generate_case_brief(prompt: str, context: str) -> str:
    """Generates a case brief based on the provided context with enhanced formatting."""
    brief_prompt = f"""
    You are an expert legal assistant. Please analyze the following legal document(s) and create a comprehensive,
    well-structured case brief. The brief should be professional, clear, and organized with the following sections:

    1. **Case Name and Citation**: Identify the full case name and any available citations.
    2. **Court and Date**: Specify the court that decided the case and the decision date.
    3. **Key Facts**: Summarize the essential facts that are relevant to the court's decision.
    4. **Procedural History**: Briefly describe how the case reached this court.
    5. **Legal Issues**: List the specific legal questions the court needed to decide.
    6. **Court's Holding**: State the court's decision on each legal issue.
    7. **Reasoning**: Explain the court's analysis and legal reasoning.
    8. **Rule of Law**: Identify the legal rule or principle established by this case.
    9. **Significance**: Explain why this case is important or how it affects existing law.
    10. **Concurring/Dissenting Opinions**: Note any significant separate opinions.

    Format the brief with clear section headers using Markdown (## for main headers, ### for subheaders).
    Use bullet points for lists and proper legal citations where applicable.

    Document Context:
    {context}

    Additional Instructions:
    {prompt if prompt else 'Provide a thorough and well-organized brief with clear sections.'}

    Case Brief:
    """
    response = get_gemini_response(brief_prompt)

    # Ensure the response is properly formatted with markdown
    if response and not response.startswith('##'):
        response = f"## Case Brief\n\n{response}"

    return response

def chunk_text(text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[str]:
    """Splits text into smaller chunks with overlap."""
    if not text:
        return []

    words = re.split(r'\s+', text)
    chunks = []

    for i in range(0, len(words), chunk_size - chunk_overlap):
        chunk = " ".join(words[i:i + chunk_size])
        if chunk.strip():  # Only add non-empty chunks
            chunks.append(chunk)

    return chunks

@st.cache_resource
def load_embedding_model():
    """Loads the sentence transformer model (cached)."""
    return SentenceTransformer('all-MiniLM-L6-v2')

def get_embeddings(text_chunks: List[str], model) -> np.ndarray:
    """Creates embeddings for text chunks and returns as numpy array."""
    embeddings = model.encode(text_chunks, convert_to_tensor=False)
    return np.array(embeddings).astype('float32')

def get_gemini_response(prompt: str) -> str:
    """Generates a response from the Gemini API with robust model fallbacks."""
    candidate_models = [
        'gemini-2.0-flash-exp',
        'gemini-2.0-flash',
        'gemini-1.5-pro',
        'gemini-1.5-flash-8b',
        'gemini-1.5-pro-latest',
        'gemini-1.0-pro'
    ]
    last_error = None
    for model_name in candidate_models:
        try:
            model = genai.GenerativeModel(model_name)
            response = model.generate_content(prompt)
            if hasattr(response, 'text') and response.text:
                return response.text
        except Exception as e:
            last_error = e
            continue
    st.error(f"An error occurred with the Gemini API. Last error: {last_error}")
    return None

# --- Streamlit App ---

# Custom Header with JurisMind branding
st.markdown(
    """
    <div style='padding: 2rem; background: linear-gradient(90deg, #831843 0%, #db2777 100%); border-radius: 15px; color: white; margin-bottom: 2rem;'>
        <div style='display: flex; align-items: center;'>
            <img src='https://img.icons8.com/color/96/000000/law.png' width='80' style='margin-right: 1.5rem;'/>
            <div>
                <h1 style='color: white; margin: 0; font-size: 2.5rem;'>JurisMind</h1>
                <p style='margin: 0.5rem 0 0 0; opacity: 0.9; font-size: 1.1rem;'>AI-Powered Legal Document Analysis with FAISS Vector Database</p>
            </div>
        </div>
    </div>
    """,
    unsafe_allow_html=True
)

# Tabs for different functionalities
tab1, tab2 = st.tabs(["Document Analysis", "Vector DB Stats"])

# --- Sidebar for Setup and Input ---
with st.sidebar:
    st.markdown("<h2 style='color: white;'> Setup</h2>", unsafe_allow_html=True)

    with st.expander("API Configuration", expanded=True):
        gemini_api_key = st.text_input("Enter your Gemini API Key:", type="password")
        if gemini_api_key:
            try:
                genai.configure(api_key=gemini_api_key)
                st.success(" API Key configured successfully!")
            except Exception as e:
                st.error(f"Error: {str(e)}")

    st.markdown("---")

    with st.expander("Input Sources", expanded=True):
        # Tabs for different input methods
        input_tab1, input_tab2 = st.tabs(["Upload", "URL"])

        with input_tab1:
            uploaded_files = st.file_uploader(
                " Upload Legal Documents",
                type=["pdf", "docx", "txt"],
                accept_multiple_files=True,
                help="Upload PDF, DOCX, or TXT files (multiple files supported)"
            )

        with input_tab2:
            url_input = st.text_input("Enter document URL", placeholder="https://example.com/legal-document")

    st.markdown("---")

    # Vector DB Configuration
    with st.expander(" Vector DB Settings", expanded=False):
        chunk_size = st.slider("Chunk Size (words)", 500, 2000, 1000, 100)
        chunk_overlap = st.slider("Chunk Overlap (words)", 50, 500, 200, 50)
        top_k_results = st.slider("Top K Results", 1, 10, 5, 1)

    st.markdown("---")

    # Quick actions
    st.markdown("### Quick Actions")
    if st.button("Clear Session", use_container_width=True):
        st.session_state.clear()
        st.rerun()

    # Help section
    with st.expander("Help"):
        st.markdown("""
        ### How to use:
        1. Upload PDFs or enter a URL
        2. Documents are processed and stored in FAISS vector DB
        3. Ask questions using semantic search
        4. Generate case briefs with one click

        ### Vector Database Benefits:
        - **Fast Similarity Search**: FAISS provides millisecond-level search
        - **Semantic Understanding**: Finds contextually relevant information
        - **Scalable**: Handles large document collections efficiently
        - **No External Dependencies**: Runs locally
        """)

# Initialize session state variables
if 'vector_db' not in st.session_state:
    st.session_state.vector_db = None
if 'messages' not in st.session_state:
    st.session_state.messages = []
if 'document_processed' not in st.session_state:
    st.session_state.document_processed = False

# Determine the current data source and its identifier
data_source = None
source_id = None
if uploaded_files:
    data_source = "files"
    source_id = ", ".join([f.name for f in uploaded_files])
elif url_input:
    data_source = "url"
    source_id = url_input

# Process data if a new source is provided
if data_source:
    # Reset session state if the source changes
    if "source_id" not in st.session_state or st.session_state.source_id != source_id:
        st.session_state.clear()
        st.session_state.source_id = source_id
        st.session_state.data_source = data_source
        st.session_state.document_processed = False
        st.session_state.messages = []
        st.session_state.vector_db = None

    if not st.session_state.document_processed:
        with st.spinner(" Processing documents and building vector database..."):
            raw_text = ""

            # Extract text based on source type
            if st.session_state.data_source == "files":
                with st.status("Processing documents...", expanded=True) as status:
                    total_files = len(uploaded_files)
                    for i, uploaded_file in enumerate(uploaded_files, 1):
                        status.write(f"Processing file {i} of {total_files}: {uploaded_file.name}")
                        file_text = extract_text_from_file(uploaded_file)
                        if file_text:
                            raw_text += file_text + "\n"
                            status.write(f"   Successfully processed {uploaded_file.name}")
                        else:
                            st.warning(f"  Could not extract text from {uploaded_file.name}")
                    status.update(label="Documents processed successfully!", state="complete", expanded=False)

            elif st.session_state.data_source == "url":
                with st.status("Processing URL content...") as status:
                    raw_text = extract_text_from_url(url_input)
                    if raw_text:
                        status.update(label="URL content processed successfully!", state="complete")
                    else:
                        status.update(label=" Could not process URL content", state="error")

            if raw_text:
                # Chunk the text
                with st.status(" Chunking text...", expanded=False) as status:
                    text_chunks = chunk_text(raw_text, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
                    if not text_chunks:
                        st.warning("Could not extract any text chunks from the source.")
                        st.stop()
                    status.update(label=f" Created {len(text_chunks)} text chunks", state="complete")

                # Load embedding model and create embeddings
                with st.status("Creating embeddings...", expanded=False) as status:
                    embedding_model = load_embedding_model()
                    embeddings = get_embeddings(text_chunks, embedding_model)
                    status.update(label=f"Generated embeddings for {len(text_chunks)} chunks", state="complete")

                # Initialize and populate FAISS vector database
                with st.status(" Building FAISS vector database...", expanded=False) as status:
                    vector_db = FAISSVectorDB(embedding_dim=embeddings.shape[1])

                    # Create metadata for each chunk
                    metadata = [{"chunk_id": i, "source": source_id} for i in range(len(text_chunks))]

                    # Add documents to vector DB
                    vector_db.add_documents(embeddings, text_chunks, metadata)

                    st.session_state.vector_db = vector_db
                    st.session_state.embedding_model = embedding_model
                    st.session_state.full_text = raw_text
                    st.session_state.document_processed = True

                    stats = vector_db.get_stats()
                    status.update(
                        label=f"Vector DB ready: {stats['total_documents']} documents indexed",
                        state="complete"
                    )

                st.success("Documents processed and indexed in FAISS vector database! You can now ask questions.")
            else:
                st.warning("Failed to get text from the source. Please check the URL or files.")
                st.stop()

# Main content area
with tab1:
    if not (uploaded_files or url_input):
        # Hero section for empty state
        col1, col2 = st.columns([2, 1])
        with col1:
            st.markdown("""
            <div class='card'>
                <h3 style='color: #831843;'>Welcome to JurisMind with FAISS</h3>
                <p>Get started by uploading legal documents or entering a URL to analyze case law using advanced vector search technology.</p>
                <div style='margin: 2rem 0;'>
                    <div style='display: flex; align-items: center; margin: 1rem 0;'>
                        <span style='background: #fce7f3; padding: 0.5rem; border-radius: 50%; margin-right: 1rem;'></span>
                        <div>
                            <h4 style='margin: 0;'>FAISS Vector Database</h4>
                            <p style='margin: 0.25rem 0 0 0; color: #64748b;'>Lightning-fast semantic search powered by Facebook's FAISS</p>
                        </div>
                    </div>
                    <div style='display: flex; align-items: center; margin: 1rem 0;'>
                        <span style='background: #fce7f3; padding: 0.5rem; border-radius: 50%; margin-right: 1rem;'></span>
                        <div>
                            <h4 style='margin: 0;'>Semantic Search</h4>
                            <p style='margin: 0.25rem 0 0 0; color: #64748b;'>Find relevant information based on meaning, not just keywords</p>
                        </div>
                    </div>
                </div>
            </div>
            """, unsafe_allow_html=True)

        with col2:
            st.markdown("""
            <div class='card' style='text-align: center; padding: 2rem 1rem;'>
                <h4>Ready to get started?</h4>
                <p>Upload your documents or enter a URL in the sidebar.</p>
                <div style='margin-top: 2rem;'>
                    <span style='font-size: 3rem;'></span>
                </div>
            </div>
            """, unsafe_allow_html=True)

    elif st.session_state.get("document_processed", False):
        # Action buttons
        col1, col2 = st.columns([1, 1])

        with col1:
            if st.button(" Generate Case Brief", use_container_width=True, type="primary"):
                with st.spinner("Analyzing documents and generating comprehensive case brief..."):
                    brief = generate_case_brief("", st.session_state.get('full_text', ''))

                    if brief:
                        formatted_brief = f"""
                        <div style='background: #fdf2f8; border-radius: 12px; border: 1px solid #fbcfe8; padding: 1.5rem; margin: 1.5rem 0; box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05);'>
                            <div style='display: flex; align-items: center; margin-bottom: 1rem; padding-bottom: 0.75rem; border-bottom: 1px solid #f9a8d4;'>
                                <span style='background: #db2777; color: white; width: 36px; height: 36px; border-radius: 50%; display: flex; align-items: center; justify-content: center; margin-right: 0.75rem;'>📋</span>
                                <h2 style='color: #831843; margin: 0;'>Case Brief</h2>
                            </div>
                            <div style='line-height: 1.7; color: #4b5563;'>
                                {brief}
                            </div>
                            <div style='margin-top: 1.5rem; padding-top: 1rem; border-top: 1px solid #f9a8d4; font-size: 0.875rem; color: #9ca3af;'>
                                Generated by JurisMind • {datetime.datetime.now().strftime("%B %d, %Y")}
                            </div>
                        </div>
                        """

                        st.session_state.messages.append({"role": "assistant", "content": formatted_brief})
                        st.rerun()
                    else:
                        st.error("Failed to generate case brief. Please try again.")

        with col2:
            if st.button(" Clear Chat", use_container_width=True):
                st.session_state.messages = []
                st.rerun()

        # Display chat messages
        for message in st.session_state.messages:
            if message["role"] == "user":
                with st.chat_message("user"):
                    st.markdown(f"<div class='user-message'>{message['content']}</div>", unsafe_allow_html=True)
            else:
                with st.chat_message("assistant"):
                    if '<div' in message['content'] or '<h' in message['content']:
                        st.markdown(message['content'], unsafe_allow_html=True)
                    else:
                        st.markdown(f"<div class='assistant-message'>{message['content']}</div>", unsafe_allow_html=True)

        # Chat input
        prompt = st.chat_input("Ask a question about the document(s)...")
        if prompt:
            st.session_state.messages.append({"role": "user", "content": prompt})

            with st.chat_message("assistant"):
                with st.spinner("Searching vector database..."):
                    # Get query embedding
                    embedding_model = st.session_state.embedding_model
                    query_embedding = get_embeddings([prompt], embedding_model)

                    # Search in FAISS vector database
                    vector_db = st.session_state.vector_db
                    relevant_chunks, distances, indices = vector_db.search(
                        query_embedding,
                        top_k=top_k_results
                    )

                if not relevant_chunks:
                    st.warning("Could not find relevant information in the document.")
                    st.stop()

                # Build context for RAG
                context = "\n\n---\n\n".join(relevant_chunks)

                rag_prompt = f"""
                Based on the following excerpts from the provided document/webpage, please answer the user's question.
                Your answer should be clear, concise, and directly based on the provided text.
                If the document doesn't contain the answer, state that clearly.

                **Context from the source:**
                {context}

                **User's Question:**
                {prompt}

                **Answer:**
                """
                
                response = get_gemini_response(rag_prompt)
                
                if response:
                    st.markdown(f"<div class='assistant-message'>{response}</div>", unsafe_allow_html=True)
                    st.session_state.messages.append({"role": "assistant", "content": response})
                else:
                    st.error("Failed to get a response from Gemini.")

with tab2:
    if st.session_state.get("vector_db"):
        st.markdown("### Vector Database Statistics")
        stats = st.session_state.vector_db.get_stats()
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Total Documents (Chunks)", stats["total_documents"])
        with col2:
            st.metric("Embedding Dimension", stats["embedding_dimension"])
        with col3:
            st.metric("Index Size", stats["index_size"])
            
        st.markdown("### Text Chunks Preview")
        st.json(st.session_state.vector_db.text_chunks[:5])
    else:
        st.info("No vector database built yet. Upload documents or provide a URL to view stats.")